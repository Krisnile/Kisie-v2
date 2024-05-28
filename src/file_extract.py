#!/usr/bin/env python3

import os
from docx import Document
import fitz  # PyMuPDF
import PyPDF2
from PIL import Image
import pytesseract

import prepare as pre
import functions as func

def extract_text_from_pdf(pdf_path):
    """
    @description  : to extract text from pdf type file
    ----------
    @param  : pdf_path(file's path)
    ----------
    @Returns  : text extracted
    ----------
    """
    
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in range(len(reader.pages)):
                text += reader.pages[page].extract_text()
    except Exception as e:
        print(f"Error extracting text from PDF '{pdf_path}': {e}")
    return text


def extract_text_from_docx(docx_path):
    """
    @description  : to extract text from docx type file
    ----------
    @param  : docx_path(file's path)
    ----------
    @Returns  : text extracted
    ----------
    """
    
    text = ''
    try:
        document = Document(docx_path) 
        col_keys = []       # Column name
        col_values = []     # Column value
        index_num = 0
        # Extract form information, add a deduplication mechanism
        fore_str = ""
        cell_text = ""
        for table in document.tables:
            for row_index,row in enumerate(table.rows):
                for col_index,cell in enumerate(row.cells):
                    if fore_str != cell.text:
                        if index_num % 2==0:
                            col_keys.append(cell.text)
                        else:
                            col_values.append(cell.text)
                        fore_str = cell.text
                        index_num +=1
                        cell_text += cell.text + '\n'
        # Extract text information
        paragraphs_text = ""
        for paragraph in document.paragraphs:
            paragraphs_text += paragraph.text + "\n"

        text = cell_text + paragraphs_text
    except Exception as e:
        print(f"Error extracting text from DOCX '{docx_path}': {e}")
    return text


def extract_text_from_png(png_path):
    """
    @description  : to extract text from png type file
    ----------
    @param  : png_path(file's path)
    ----------
    @Returns  : text extracted
    ----------
    """
   
    # Open image
    img = Image.open(png_path)

    # Character recognition
    text = pytesseract.image_to_string(img, lang = 'chi_sim+eng_sim')

    text = ''.join([i.strip(' ') for i in text])
    # Print results
    print("Result of Extract: ", text)

    return text


def save_text_to_file(text_list, file_path):
    """
    @description  : save text to a given path
    ----------
    @param  : text file_path
    ----------
    @Returns  : None
    ----------
    """
    
    try:
        for text in text_list:
            with open(file_path, 'a', encoding='utf-8') as file:
                file.write(text)
        print(f"Text saved to '{file_path}'\n")
    except Exception as e:
        print(f"Error saving text to file '{file_path}': {e}")


def batch_extract_text_from_files(path_lists):
    """
    @description  : batch extract text from files
    ----------
    @param  : path_lists(a list of file waiting to extract text)
    ----------
    @Returns  : text_list(a list of text)
    ----------
    """
    try:
        text_list = list()
        for file_path in path_lists:
            if file_path.endswith('.pdf'):
                text_list.append( extract_text_from_pdf(file_path))
            elif file_path.endswith('.docx'):
                text_list.append(extract_text_from_docx(file_path))
            elif file_path.endswith('.png'):
                text_list.append(extract_text_from_png(file_path))
            else:
                file_name = os.path.basename(file_path)
                print(f"unsupported formats: {file_name}")
                print("ignore the file\n")
                continue  # Ignore files with unsupported formats
            
        # Save extracted text to a file with the same name
        file_path_saved = pre.File_Extract_Save
        text_file_path = (f'{file_path_saved}/chatbot_dialogues_{func.time}.txt')
        save_text_to_file(text_list, text_file_path)

        return text_list
    
    except Exception as e:
        print(f"Error extracting text from file '{path_lists}': {e}")